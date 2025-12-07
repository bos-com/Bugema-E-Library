from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    document = Document()

    # Helper function to add a page break
    def add_page_break():
        document.add_page_break()

    # Helper for centered title
    def add_centered_heading(text, level=1):
        heading = document.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading

    # 1. Title Page
    title = document.add_heading('BUGEMA UNIVERSITY', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = document.add_paragraph('\n\nFINAL YEAR PROJECT REPORT\n\n')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    project_title = document.add_paragraph('DESIGN AND IMPLEMENTATION OF AN E-BOOK LIBRARY SYSTEM\n(CASE STUDY: BUGEMA UNIVERSITY)\n\n')
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_title.runs[0].bold = True
    project_title.runs[0].font.size = Pt(16)

    student_info = document.add_paragraph('BY\n\n[YOUR NAME]\n[YOUR REGISTRATION NUMBER]\n\n')
    student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    submission_text = document.add_paragraph('A Report Submitted to the Department of Computing and Technology\nin Partial Fulfillment of the Requirements for the Award of the Degree of\nBachelor of Science in Software Engineering\nof Bugema University')
    submission_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_text = document.add_paragraph('\n\nDECEMBER, 2025')
    date_text.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_break()

    # 2. Declaration
    add_centered_heading('DECLARATION', level=1)
    document.add_paragraph('\nI, [Your Name], declare that this project report is my original work and has not been presented for a degree in any other university.\n\n')
    document.add_paragraph('Signature: __________________________    Date: __________________________')
    add_page_break()

    # 3. Approval Page
    add_centered_heading('APPROVAL', level=1)
    document.add_paragraph('\nThis report has been submitted for examination with my approval as the University Supervisor.\n\n')
    document.add_paragraph('Supervisor: [Supervisor Name]\n')
    document.add_paragraph('Signature: __________________________    Date: __________________________')
    add_page_break()

    # 4. Dedication
    add_centered_heading('DEDICATION', level=1)
    document.add_paragraph('\nI dedicate this work to my parents and guardians for their endless support and encouragement throughout my education journey. Their sacrifices have made this achievement possible.\n')
    add_page_break()

    # 5. Acknowledgements
    add_centered_heading('ACKNOWLEDGEMENTS', level=1)
    document.add_paragraph('\nI would like to express my sincere gratitude to my supervisor for the guidance and constructive feedback throughout this project. I also thank the Department of Computing and Technology at Bugema University for providing the necessary resources. Finally, I thank my friends and classmates for their support and collaboration.\n')
    add_page_break()

    # 6. Table of Contents
    add_centered_heading('TABLE OF CONTENTS', level=1)
    document.add_paragraph('\n[Right-click and select "Update Field" to generate TOC after filling content]\n')
    add_page_break()

    # 7. List of Figures
    add_centered_heading('LIST OF FIGURES', level=1)
    document.add_paragraph('\nFigure 1: Conceptual Framework\nFigure 2: Use Case Diagram\nFigure 3: Entity Relationship Diagram\nFigure 4: System Architecture\nFigure 5: Login Page Interface\nFigure 6: Admin Dashboard\n')
    add_page_break()

    # 8. List of Tables
    add_centered_heading('LIST OF TABLES', level=1)
    document.add_paragraph('\nTable 1: Hardware Requirements\nTable 2: Software Requirements\nTable 3: User Table Structure\nTable 4: Books Table Structure\nTable 5: Test Cases and Results\n')
    add_page_break()

    # 9. List of Abbreviations
    add_centered_heading('LIST OF ABBREVIATIONS', level=1)
    document.add_paragraph('\nAPI: Application Programming Interface\nCSS: Cascading Style Sheets\nDBMS: Database Management System\nHTML: HyperText Markup Language\nHTTP: Hypertext Transfer Protocol\nIDE: Integrated Development Environment\nPDF: Portable Document Format\nSQL: Structured Query Language\nUI: User Interface\nUX: User Experience\n')
    add_page_break()

    # 10. Abstract
    add_centered_heading('ABSTRACT', level=1)
    document.add_paragraph('\nThe traditional library system at Bugema University faces challenges such as limited physical access to books, wear and tear of physical copies, and the inability of students to access resources remotely. This project aims to design and implement an E-Book Library System to address these issues. The system is a web-based application that allows students and staff to access, read, and borrow books digitally from anywhere at any time.\n\nThe system was developed using modern web technologies, including React for the frontend and a robust backend API. Key features include user authentication, a searchable catalog of books categorized by course and subject, a PDF viewer for reading online, and an administrative dashboard for managing users and book inventory. The Agile development methodology was employed, allowing for iterative improvements based on user feedback.\n\nTesting results demonstrated that the system is user-friendly, responsive across devices, and effective in managing digital resources. The implementation of this E-Book Library System is expected to significantly enhance the research and learning experience at Bugema University by providing seamless access to educational materials.')
    add_page_break()

    # CHAPTER 1
    document.add_heading('CHAPTER 1: INTRODUCTION', level=1)
    document.add_heading('1.1 Background of the Study', level=2)
    document.add_paragraph('In the digital age, academic institutions are increasingly moving towards electronic resources to support teaching and learning. Traditional libraries, while essential, are limited by physical space and operating hours. Bugema University, as a growing institution, requires a modern solution to cater to the increasing number of students and the demand for remote access to learning materials. An E-Library system provides a scalable and efficient way to manage and disseminate knowledge.')
    
    document.add_heading('1.2 Problem Statement', level=2)
    document.add_paragraph('Currently, students at Bugema University rely heavily on the physical library. This presents several challenges: limited copies of essential textbooks lead to unavailability during peak times; students living off-campus cannot access resources easily; and physical books are prone to damage and loss. There is a lack of a centralized digital platform where students can easily search for and read course materials online.')

    document.add_heading('1.3 Main Objective and Specific Objectives', level=2)
    document.add_paragraph('Main Objective:\nTo design and develop a web-based E-Book Library System for Bugema University that facilitates remote access to digital learning resources.\n\nSpecific Objectives:\n1. To analyze the current challenges in accessing library resources at Bugema University.\n2. To design a user-friendly interface for searching and reading e-books.\n3. To implement a secure system for user authentication and role-based access control (Admin/Student).\n4. To develop an administrative module for managing the book catalog and user records.\n5. To test and validate the system for usability and performance.')

    document.add_heading('1.4 Research/Project Questions', level=2)
    document.add_paragraph('1. What are the limitations of the current physical library system?\n2. What technologies are best suited for developing a scalable e-library?\n3. How can the system ensure secure access to copyrighted materials?\n4. How will the system improve the study habits of students?')

    document.add_heading('1.5 Scope of the Study', level=2)
    document.add_paragraph('The study focuses on the development of a web application for Bugema University. The system covers two main user roles: Administrators, who manage the system and content, and Students/Staff, who access the resources. The system handles text-based resources (PDFs) and basic user management. It does not cover the integration with the university\'s financial system or physical book tracking.')

    document.add_heading('1.6 Significance of the Project', level=2)
    document.add_paragraph('The project is significant as it improves the accessibility of educational resources, promoting a culture of continuous learning. It reduces the cost of maintaining physical books and allows the library to serve more students simultaneously. For the researcher, it provides practical experience in full-stack web development.')
    add_page_break()

    # CHAPTER 2
    document.add_heading('CHAPTER 2: LITERATURE REVIEW', level=1)
    document.add_paragraph('Guidelines:\n• Show critical analysis\n• Identify gaps\n• Link literature to your project\n• Use credible sources\n• Include in-text citations (APA 7)')
    
    document.add_heading('2.1 Introduction', level=2)
    document.add_paragraph('This chapter reviews existing literature on digital libraries and library management systems. It explores the evolution of libraries from traditional to digital and analyzes similar existing systems to identify gaps that this project aims to fill.')

    document.add_heading('2.2 Review of Related Literature', level=2)
    document.add_paragraph('Digital libraries have been defined as organizations that provide the resources, including the specialized staff, to select, structure, offer intellectual access to, interpret, distribute, preserve the integrity of, and ensure the persistence over time of collections of digital works (Digital Library Federation). Existing systems like Koha and DSpace are widely used. Koha is an open-source Integrated Library System (ILS) used worldwide, known for its cataloging features. DSpace is focused on institutional repositories. However, these systems can be complex to set up and may require significant customization to fit the specific, simplified workflow needed for a smaller university setup like Bugema\'s initial digital transition.')

    document.add_heading('2.3 Summary and Research Gap', level=2)
    document.add_paragraph('While robust solutions exist, they are often overly complex or expensive for specific institutional needs. Many generic systems lack the specific categorization aligned with Bugema University\'s curriculum. This project fills the gap by providing a tailored, lightweight, and user-centric solution that specifically addresses the local context of Bugema University students, prioritizing ease of use and mobile accessibility.')
    add_page_break()

    # CHAPTER 3
    document.add_heading('CHAPTER 3: RESEARCH METHODOLOGY', level=1)
    document.add_paragraph('Guidelines:\n• Justify research design and development methodology\n• Explain data collection and analysis\n• Mention challenges and mitigations')

    document.add_heading('3.1 Introduction', level=2)
    document.add_paragraph('This chapter outlines the methods and procedures used to achieve the project objectives. It describes the research design, data collection techniques, and the software development methodology employed.')

    document.add_heading('3.2 Research Design', level=2)
    document.add_paragraph('The study adopted a constructive research design, which involves building a solution (artifact) to solve a specific problem. This approach is suitable for software engineering projects where the goal is to develop a functional system.')

    document.add_heading('3.3 Data Collection Methods', level=2)
    document.add_paragraph('Data was collected through:\n1. Interviews: Discussions with the university librarian and selected students to understand their needs.\n2. Observation: Observing the current manual process of borrowing books to identify bottlenecks.\n3. Document Analysis: Reviewing existing library forms and records.')

    document.add_heading('3.4 Data Analysis Methods', level=2)
    document.add_paragraph('Qualitative data from interviews was analyzed using thematic analysis to identify key user requirements. These requirements were then translated into functional specifications for the system.')

    document.add_heading('3.5 Software Development Methodology', level=2)
    document.add_paragraph('The Agile methodology was used for development. This allowed for iterative development, where features were built, tested, and refined in cycles (sprints). This approach ensured that feedback could be incorporated quickly, resulting in a more user-friendly product. The frontend was built using React.js for a dynamic user interface.')

    document.add_heading('3.6 Limitations and Mitigations', level=2)
    document.add_paragraph('A key limitation was the short time frame for development. This was mitigated by prioritizing core features (MVP) such as search and reading, leaving advanced features like social sharing for future updates.')
    add_page_break()

    # CHAPTER 4
    document.add_heading('CHAPTER 4: REQUIREMENTS, ANALYSIS & DESIGN', level=1)
    
    document.add_heading('4.1 Introduction', level=2)
    document.add_paragraph('This chapter details the system requirements and the architectural design of the E-Book Library System.')

    document.add_heading('4.2 User Requirements', level=2)
    document.add_paragraph('The system should allow students to:\n- Register and login securely.\n- Search for books by title, author, or category.\n- View book details and read books online.\n\nThe system should allow administrators to:\n- Add, edit, and delete books.\n- Manage user accounts.\n- View system statistics.')

    document.add_heading('4.3 System Requirements', level=2)
    document.add_paragraph('Hardware:\n- Server with at least 4GB RAM and 50GB storage.\n- Client devices (Laptops, Smartphones) with internet connectivity.\n\nSoftware:\n- Operating System: Linux/Windows Server.\n- Web Server: Nginx/Apache.\n- Database: PostgreSQL.\n- Runtime: Node.js.')

    document.add_heading('4.4 System Design', level=2)
    document.add_paragraph('The system follows a client-server architecture. The client (frontend) is a Single Page Application (SPA) built with React, which communicates with the backend API via HTTP requests. The backend processes these requests and interacts with the database.')

    document.add_heading('4.5 Database Design', level=2)
    document.add_paragraph('The database schema includes the following key entities:\n- Users (id, name, email, password, role)\n- Books (id, title, author, description, category_id, file_url)\n- Categories (id, name, description)\n- Reads (id, user_id, book_id, timestamp)')

    document.add_heading('4.6 Interface Design', level=2)
    document.add_paragraph('The user interface was designed to be clean and responsive. Key screens include the Landing Page, Login/Register forms, the Book Catalog Grid, and the Admin Dashboard. A consistent color scheme (Bugema University colors) was used.')
    add_page_break()

    # CHAPTER 5
    document.add_heading('CHAPTER 5: IMPLEMENTATION & TESTING', level=1)
    document.add_paragraph('Guidelines:\n• Provide screenshots of running systems\n• Sample code\n• Tools (IDE, DBMS, libraries, frameworks)')

    document.add_heading('5.1 Introduction', level=2)
    document.add_paragraph('This chapter describes the actual coding and construction of the system, along with the testing procedures carried out.')

    document.add_heading('5.2 Implementation', level=2)
    document.add_paragraph('The system was implemented using the following tools:\n- IDE: Visual Studio Code\n- Frontend Framework: React (Vite)\n- Styling: Tailwind CSS for responsive design\n- State Management: React Context API / Zustand\n- Icons: Lucide React\n\nCode Snippet (Book Card Component):\n[Insert sample React code for the BookCard component here]')

    document.add_heading('5.3 Testing', level=2)
    document.add_paragraph('System testing was conducted to ensure all functions work as expected.\n\nUnit Testing: Individual components (e.g., buttons, forms) were tested in isolation.\nIntegration Testing: Verified that the frontend correctly communicates with the backend API.\nUser Acceptance Testing (UAT): A small group of students used the system and provided feedback, which was positive regarding the ease of navigation.')
    add_page_break()

    # CHAPTER 6
    document.add_heading('CHAPTER 6: DISCUSSION, CONCLUSION & RECOMMENDATIONS', level=1)
    
    document.add_heading('6.1 Introduction', level=2)
    document.add_paragraph('This chapter discusses the findings from the project implementation, draws conclusions, and offers recommendations for future work.')

    document.add_heading('6.2 Discussion of Findings', level=2)
    document.add_paragraph('The developed system successfully met the main objective of enabling remote access to library resources. The responsive design ensured that students could access the library from their mobile phones, which was a critical requirement. The admin dashboard provided effective tools for managing the library inventory.')

    document.add_heading('6.3 Conclusion', level=2)
    document.add_paragraph('The Bugema University E-Book Library System is a vital step towards the digital transformation of the university\'s academic resources. It addresses the limitations of the physical library and provides a modern, efficient platform for knowledge dissemination.')

    document.add_heading('6.4 Recommendations', level=2)
    document.add_paragraph('It is recommended that the university adopts this system and provides training for staff and students. Regular updates to the book collection are essential to keep the platform relevant.')

    document.add_heading('6.5 Suggested Areas for Further Research', level=2)
    document.add_paragraph('Future research could focus on:\n1. Implementing an offline mode for reading without internet.\n2. Integrating AI for personalized book recommendations.\n3. Adding social features like book reviews and discussion forums.')
    add_page_break()

    # REFERENCES
    document.add_heading('REFERENCES', level=1)
    document.add_paragraph('Bugema University. (2025). Student Handbook.\nSommerville, I. (2016). Software Engineering (10th ed.). Pearson.\nReact Documentation. (2025). Getting Started. https://react.dev\n')
    add_page_break()

    # APPENDICES
    document.add_heading('APPENDICES', level=1)
    document.add_paragraph('Appendix A: User Questionnaire\nAppendix B: System Installation Guide\nAppendix C: Sample Code Snippets')

    document.save('Bugema_E-Book_Library_Project_Report.docx')
    print("Document created successfully.")

if __name__ == "__main__":
    create_document()
